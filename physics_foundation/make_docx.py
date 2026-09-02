"""
make_docx.py
------------
Render CHF_Physics_Foundation.md into a .docx formatted to the house style
declared at the bottom of `outline/paper outline.docx`:

    "Texts should be times new roman, 12, 1.5 spacing"

Handles the markdown subset actually used in the foundation document:
headings, paragraphs, bullet/numbered lists, blockquotes, fenced code blocks
(equations), pipe tables, horizontal rules, and inline **bold** / `code`.

Run:  python physics_foundation/make_docx.py
"""
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
SRC = HERE / "CHF_Physics_Foundation.md"
DST = HERE / "CHF_Physics_Foundation.docx"

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(12)
# Equations/code stay monospace or the aligned ASCII maths collapses.
MONO_FONT = "Consolas"
MONO_SIZE = Pt(9)
LINE_SPACING = 1.5


def style_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    normal.paragraph_format.line_spacing = LINE_SPACING
    normal.paragraph_format.space_after = Pt(6)
    # Headings inherit Times New Roman too, per the outline's figure/text spec.
    for level in range(1, 5):
        st = doc.styles[f"Heading {level}"]
        st.font.name = BODY_FONT
        st.font.color.rgb = RGBColor(0, 0, 0)
        st.font.bold = True
        st.paragraph_format.line_spacing = LINE_SPACING


INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`)")


def add_inline(par, text):
    """Emit text into `par`, honouring **bold** and `code` spans."""
    for chunk in INLINE.split(text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            run = par.add_run(chunk[2:-2])
            run.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            run = par.add_run(chunk[1:-1])
            run.font.name = MONO_FONT
            run.font.size = Pt(10.5)
        else:
            par.add_run(chunk)


def add_code_block(doc, lines):
    par = doc.add_paragraph()
    pf = par.paragraph_format
    pf.line_spacing = 1.0
    pf.space_after = Pt(10)
    pf.space_before = Pt(6)
    pf.left_indent = Pt(18)
    run = par.add_run("\n".join(lines))
    run.font.name = MONO_FONT
    run.font.size = MONO_SIZE


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    table = doc.add_table(rows=len(rows), cols=len(header))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, cell_text in enumerate(header):
        cell = table.cell(0, j)
        cell.text = ""
        par = cell.paragraphs[0]
        par.paragraph_format.line_spacing = 1.0
        par.paragraph_format.space_after = Pt(2)
        add_inline(par, cell_text)
        for run in par.runs:
            run.bold = True
            run.font.size = Pt(10)
    for i, row in enumerate(body, start=1):
        for j in range(len(header)):
            cell = table.cell(i, j)
            cell.text = ""
            par = cell.paragraphs[0]
            par.paragraph_format.line_spacing = 1.0
            par.paragraph_format.space_after = Pt(2)
            add_inline(par, row[j] if j < len(row) else "")
            for run in par.runs:
                run.font.size = Pt(10)
    doc.add_paragraph()


def convert(md_text, doc):
    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # fenced code block -> equation block
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            add_code_block(doc, buf)
            continue

        # pipe table
        if stripped.startswith("|") and i + 1 < len(lines) and set(
            lines[i + 1].strip().replace("|", "").replace(" ", "")
        ) <= {"-", ":"} and lines[i + 1].strip().startswith("|"):
            rows = [split_row(stripped)]
            i += 2
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            add_table(doc, rows)
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("---"):
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = par.add_run("• • •")
            run.font.size = Pt(10)
            i += 1
            continue

        m = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            par = doc.add_heading(level=level)
            add_inline(par, m.group(2))
            for run in par.runs:
                run.font.name = BODY_FONT
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        if stripped.startswith(">"):
            buf = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(24)
            par.paragraph_format.right_indent = Pt(24)
            add_inline(par, " ".join(b for b in buf if b))
            for run in par.runs:
                run.italic = True
            continue

        m = re.match(r"^[-*]\s+(.*)", stripped)
        if m:
            par = doc.add_paragraph(style="List Bullet")
            par.paragraph_format.line_spacing = LINE_SPACING
            add_inline(par, m.group(1))
            i += 1
            continue

        m = re.match(r"^\d+\.\s+(.*)", stripped)
        if m:
            par = doc.add_paragraph(style="List Number")
            par.paragraph_format.line_spacing = LINE_SPACING
            add_inline(par, m.group(1))
            i += 1
            continue

        par = doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(par, stripped)
        i += 1


def main():
    doc = Document()
    style_document(doc)
    convert(SRC.read_text(encoding="utf-8"), doc)
    doc.save(DST)
    print(f"wrote {DST}")


if __name__ == "__main__":
    main()
